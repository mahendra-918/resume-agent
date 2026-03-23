from __future__ import annotations

import asyncio
import re
from pathlib import Path

from loguru import logger

from resume_agent.core.config import settings
from resume_agent.core.exceptions import ResumeGenerationError
from resume_agent.core.models import Job, TailoredResume
from resume_agent.core.state import AgentState


def _sanitize(text: str) -> str:
    return re.sub(r"[^\w]", "_", text).strip("_")


async def _generate_pdf(tailored: TailoredResume, output_path: Path) -> None:
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML

    template_dir = Path(__file__).parent.parent / "document" / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    template = env.get_template("resume.html")

    html_content = template.render(resume=tailored)

    def _write() -> None:
        HTML(string=html_content).write_pdf(str(output_path))

    await asyncio.to_thread(_write)
    logger.info(f"[ResumeGenerator] PDF written: {output_path}")


async def _generate_docx(tailored: TailoredResume, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    def _build() -> None:
        doc = Document()
        base = tailored.base

        # Name as title
        doc.add_heading(base.name, level=0)

        # Contact info
        contact_parts = [base.email]
        if base.phone:
            contact_parts.append(base.phone)
        if base.location:
            contact_parts.append(base.location)
        if base.linkedin:
            contact_parts.append(base.linkedin)
        if base.github:
            contact_parts.append(base.github)
        if base.portfolio:
            contact_parts.append(base.portfolio)
        doc.add_paragraph(" | ".join(contact_parts))

        # Summary
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(tailored.tailored_summary)

        # Skills
        doc.add_heading("Skills", level=1)
        skills = base.skills
        skill_groups = {
            "Languages": skills.languages,
            "Frameworks": skills.frameworks,
            "ML/AI": skills.ml_ai,
            "Cloud/DevOps": skills.cloud_devops,
            "Databases": skills.databases,
            "Tools": skills.tools,
        }
        for category, items in skill_groups.items():
            if items:
                doc.add_paragraph(f"{category}: {', '.join(items)}")

        # Experience
        if tailored.reordered_experience:
            doc.add_heading("Experience", level=1)
            for exp in tailored.reordered_experience:
                doc.add_heading(f"{exp.title} — {exp.org}", level=2)
                doc.add_paragraph(exp.duration)
                for h in exp.highlights:
                    doc.add_paragraph(h, style="List Bullet")

        # Projects
        if tailored.reordered_projects:
            doc.add_heading("Projects", level=1)
            for proj in tailored.reordered_projects:
                doc.add_heading(proj.name, level=2)
                if proj.tech_stack:
                    doc.add_paragraph(f"Tech: {', '.join(proj.tech_stack)}")
                doc.add_paragraph(proj.description)
                for h in proj.highlights:
                    doc.add_paragraph(h, style="List Bullet")

        # Education
        if base.education:
            doc.add_heading("Education", level=1)
            for edu in base.education:
                p = doc.add_paragraph()
                p.add_run(f"{edu.degree} — {edu.institution}").bold = True
                doc.add_paragraph(edu.duration + (f" | GPA: {edu.gpa}" if edu.gpa else ""))

        doc.save(str(output_path))

    await asyncio.to_thread(_build)
    logger.info(f"[ResumeGenerator] DOCX written: {output_path}")


async def generate_resume_node(state: AgentState) -> dict:
    tailored: TailoredResume = state["tailored_resume"]
    job: Job = state["current_job"]

    output_dir = Path(settings.OUTPUT_DIR)
    await asyncio.to_thread(output_dir.mkdir, parents=True, exist_ok=True)

    filename = f"{_sanitize(job.company)}_{_sanitize(job.title)}_resume"
    pdf_path = output_dir / f"{filename}.pdf"
    docx_path = output_dir / f"{filename}.docx"

    try:
        await _generate_pdf(tailored, pdf_path)
        tailored.pdf_path = str(pdf_path)
    except Exception as e:
        logger.error(f"[ResumeGenerator] PDF generation failed: {e}")
        errors = list(state.get("errors") or [])
        errors.append(f"ResumeGenerator PDF: {e}")
        return {"tailored_resume": tailored, "errors": errors}

    try:
        await _generate_docx(tailored, docx_path)
        tailored.docx_path = str(docx_path)
    except Exception as e:
        logger.error(f"[ResumeGenerator] DOCX generation failed: {e}")
        errors = list(state.get("errors") or [])
        errors.append(f"ResumeGenerator DOCX: {e}")

    return {"tailored_resume": tailored}
